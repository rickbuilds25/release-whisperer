import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TPG_RED = RGBColor(244, 95, 95)

def format_to_docx(data, path):
    doc = Document()

    # Title
    title = doc.add_heading("TPG Release Summary", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = TPG_RED

    for entry in data:
        doc.add_heading(f"🔹 {entry['ticket_id']} — {entry['summary']}", level=2)
        for tone, text in entry["summaries"].items():
            para = doc.add_paragraph()
            para.add_run(f"🧠 {tone.title()} Summary:\n").bold = True
            para.add_run(text + "\n")
        doc.add_paragraph("—" * 40)

    # Footer
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.text = "Made with 💕 by THE PRODUCT GEEK  |  theproductgeek.club  |  Instagram: @the.productgeek"
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.runs[0].font.size = Pt(9)

    doc.save(path)
    print(f"✅ .docx saved to {path}")


def format_to_markdown(data, path):
    lines = ["# 📦 TPG Release Summary\n"]
    for entry in data:
        lines.append(f"## 🔹 {entry['ticket_id']} — {entry['summary']}")
        for tone, text in entry["summaries"].items():
            lines.append(f"**🧠 {tone.title()} Summary:**\n{text}\n")
        lines.append("---")
    with open(path, "w") as f:
        f.write("\n".join(lines))
    print(f"✅ Markdown saved to {path}")


def format_to_slack_block(data, path):
    lines = ["*TPG Release Summary*"]
    for entry in data:
        lines.append(f"\n• *{entry['ticket_id']} – {entry['summary']}*")
        for tone, text in entry["summaries"].items():
            lines.append(f"> *{tone.title()} Summary*: {text}")
    with open(path, "w") as f:
        f.write("\n".join(lines))
    print(f"✅ Slack block saved to {path}")
